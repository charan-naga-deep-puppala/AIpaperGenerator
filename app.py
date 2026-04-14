import os
import io
from flask import Flask, render_template, request, send_file, jsonify
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Configure Gemini
GENAI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GENAI_API_KEY)

def generate_paper_content(subject, grade, time, marks, prompt_text):
    """Sends the prompt to Gemini to get the questions."""
    
    model = genai.GenerativeModel('gemini-2.5-flash')
    
    # We instruct Gemini to just give us the questions/sections cleanly
    full_prompt = f"""
    Create a formal exam question paper for:
    Subject: {subject}
    Class: {grade}
    Time: {time}
    Max Marks: {marks}
    
    Specific Instructions: {prompt_text}
    
    Format the output clearly with Section headings (Section A, Section B) and numbered questions. 
    Do NOT include the school header or title in your output (I will add that manually). 
    Just start with the instructions or the first section.
    """
    
    response = model.generate_content(full_prompt)
    return response.text

def create_word_doc(header_info, content):
    """Creates a formatted .docx file with the specific school header."""
    doc = Document()
    
    # --- 1. THE SPECIFIC HEADER (Based on your image) ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Line 1: School Name
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("BHARATIYA VIDYA BHAVAN’S RESIDENTIAL PUBLIC SCHOOL")
    run1.bold = True
    run1.font.size = Pt(14)

    # Line 2: Sub-header
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("(Sponsored by rice millers’ education and cultural society, W.G.Dt)")
    run2.font.size = Pt(10)

    # Line 3: Location
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("VIDHYASHRAM – BHIMAVARAM-534201")
    run3.bold = True
    run3.font.size = Pt(12)

    # Line 4: Test Name
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"{header_info['test_name']} – {header_info['academic_year']}")
    run4.bold = True
    run4.font.size = Pt(12)

    doc.add_paragraph() # Spacer

    # Line 5: Details (Subject/Time/Class/Marks)
    # We use a table for alignment (Left vs Right)
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    
    # Row 1
    table.cell(0, 0).text = f"Sub: {header_info['subject']}"
    cell_right = table.cell(0, 1)
    cell_right.text = f"Time: {header_info['time']}"
    cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Row 2
    table.cell(1, 0).text = f"Class: {header_info['grade']}"
    cell_right_2 = table.cell(1, 1)
    cell_right_2.text = f"Marks: {header_info['marks']}"
    cell_right_2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Bold the table text
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    # Horizontal Line (visualized using underscores for simplicity in docx)
    p_line = doc.add_paragraph("_" * 80)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer

    # --- 2. THE GENERATED QUESTIONS ---
    # We simply dump the Gemini text here. 
    # For a production app, you might parse Markdown, but text is fine for now.
    doc.add_paragraph(content)

    # Save to memory buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    data = request.json
    
    # Extract data from UI
    subject = data.get('subject', 'Business Studies')
    grade = data.get('grade', 'XI IEC & PEC')
    test_name = data.get('testName', 'Revision Test - XI')
    year = data.get('year', '2025-26')
    time = data.get('time', '40min')
    marks = data.get('marks', '20')
    user_prompt = data.get('prompt', '')

    # 1. Call Gemini
    try:
        generated_text = generate_paper_content(subject, grade, time, marks, user_prompt)
        
        # Return the text so the frontend can preview it
        return jsonify({'success': True, 'content': generated_text})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/download', methods=['POST'])
def download():
    # Helper to generate the doc when the user clicks "Download"
    # Note: In a real app, you might cache the last generation. 
    # Here, we will accept the content back from the frontend or regenerate quickly.
    # To keep it simple, let's assume we regenerate or pass content hidden.
    
    # BETTER APPROACH: The Frontend sends the header data AND the content (that was already generated)
    data = request.form
    
    header_info = {
        'subject': data.get('subject'),
        'grade': data.get('grade'),
        'test_name': data.get('testName'),
        'academic_year': data.get('year'),
        'time': data.get('time'),
        'marks': data.get('marks')
    }
    content = data.get('content') # The text Gemini generated
    
    file_stream = create_word_doc(header_info, content)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"{header_info['subject']}_Test.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)